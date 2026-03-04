import streamlit as st
import os
import base64
import sys
from typing import TypedDict, List
from pypdf import PdfReader

if sys.stdout.encoding != 'utf-8':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from langgraph.graph import StateGraph, END
from langchain_groq import ChatGroq
from langchain_core.messages import BaseMessage, HumanMessage, SystemMessage, AIMessage
from groq import Groq
from docx import Document as _Document
from docx.shared import Pt as _Pt, RGBColor as _RGBColor
import io as _io

# ─────────────────────────────────────────────
# DOCUMENTOS DE REFERENCIA PARA EL MODO DOCENTE
# ─────────────────────────────────────────────
# Para agregar un nuevo documento:
#   1. Poné el PDF en la carpeta  docs/
#   2. Agregá una entrada acá con:
#      "Nombre que aparece en el checkbox": {
#          "archivo": "nombre_del_archivo.pdf",
#          "emoji":   "📋",
#          "descripcion": "Texto corto que se muestra en el sidebar",
#      }
DOCS_CONFIG = {
    "NRA 2026": {
        "archivo":     "nra.pdf",
        "emoji":       "📋",
        "descripcion": "Nuevo Régimen Académico · Ministerio de Educación de Córdoba",
    },
    "Marco Curricular Común": {
        "archivo":     "MARCO CURRICULAR COMUN.pdf",
        "emoji":       "📘",
        "descripcion": "Marco Curricular Común",
    },
    "Orientaciones": {
        "archivo":     "ORIENTACIONES.pdf",
        "emoji":       "🧭",
        "descripcion": "Orientaciones generales",
    },
    "Tecnicaturas": {
        "archivo":     "TECNICATURAS.pdf",
        "emoji":       "🔧",
        "descripcion": "Tecnicaturas",
    },
    "Progresiones Formación Vida y Trabajo": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE FORMACION PARA LA VIDA Y EL TRABAJO.pdf",
        "emoji":       "💼",
        "descripcion": "Progresiones de Aprendizaje · Formación para la Vida y el Trabajo",
    },
    "Progresiones Educación Física": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE EDUCACION FISICA.pdf",
        "emoji":       "⚽",
        "descripcion": "Progresiones de Aprendizaje · Educación Física",
    },
    "Progresiones Ed. Tecnológica": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE EDUCACION TECNOLOGICA Y CIENCIAS DE LA COMPUTACION.pdf",
        "emoji":       "💻",
        "descripcion": "Progresiones de Aprendizaje · Educación Tecnológica y Ciencias",
    },
    "Progresiones Ed. Artística": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE EDUCACION ARTISTICA.pdf",
        "emoji":       "🎨",
        "descripcion": "Progresiones de Aprendizaje · Educación Artística",
    },
    "Progresiones Matemática": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE MATEMATICA.pdf",
        "emoji":       "🔢",
        "descripcion": "Progresiones de Aprendizaje · Matemática",
    },
    "Progresiones Lengua y Literatura": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE LENGUAJE-LENGUA Y LITERATURA.pdf",
        "emoji":       "📖",
        "descripcion": "Progresiones de Aprendizaje · Lenguaje, Lengua y Literatura",
    },
    "Progresiones Cs. Naturales": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE CS NATURALES.pdf",
        "emoji":       "🔬",
        "descripcion": "Progresiones de Aprendizaje · Ciencias Naturales",
    },
    "Progresiones Cs. Sociales": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE CS SOCIALES.pdf",
        "emoji":       "🌍",
        "descripcion": "Progresiones de Aprendizaje · Ciencias Sociales",
    },
    "Progresiones Ciudadanía": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE CIUDADANIA Y HUMANIDADES.pdf",
        "emoji":       "🏛️",
        "descripcion": "Progresiones de Aprendizaje · Ciudadanía y Humanidades",
    },
    "Progresiones Inglés": {
        "archivo":     "PROGRESIONES DE APRENDIZAJE DE LENGUA EXTRANJERA INGLES.pdf",
        "emoji":       "🇬🇧",
        "descripcion": "Progresiones de Aprendizaje · Lengua Extranjera Inglés",
    },
    "Orient. Pedagógicas Inicial": {
        "archivo":     "ORIENTACIONES PEDAGOGICAS Y DIDACTICAS  EDUCACION INICIAL.pdf",
        "emoji":       "📗",
        "descripcion": "Orientaciones Pedagógicas y Didácticas · Educación Inicial",
    },
    "Orient. Pedagógicas Primaria": {
        "archivo":     "EDUCACION PRIMARIA - ORIENTACIONES PEDAGOGICAS Y DIDACTICAS.pdf",
        "emoji":       "📙",
        "descripcion": "Orientaciones Pedagógicas y Didácticas · Educación Primaria",
    },
    "Orient. Pedagógicas Secundaria": {
        "archivo":     "EDUCACION SECUNDARIA - ORIENTACIONES PEDAGOGICAS Y DIDACTICAS.pdf",
        "emoji":       "📕",
        "descripcion": "Orientaciones Pedagógicas y Didácticas · Educación Secundaria",
    },
}

import unicodedata as _unicodedata

def _normalizar_nombre(texto: str) -> str:
    """Elimina tildes y pasa a mayúsculas para comparación tolerante."""
    return _unicodedata.normalize('NFD', texto).encode('ascii', 'ignore').decode('utf-8').upper()

@st.cache_resource
def cargar_documentos_referencia() -> dict:
    base = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")
    resultado = {}

    if not os.path.exists(base):
        return resultado

    # Mapa de nombre_normalizado → nombre_real para buscar sin importar tildes
    archivos_en_disco = {}
    for f in os.listdir(base):
        archivos_en_disco[_normalizar_nombre(f)] = f

    for nombre, cfg in DOCS_CONFIG.items():
        archivo = cfg["archivo"]

        # Buscar primero versión .txt (PDFs escaneados pre-convertidos)
        nombre_txt = _normalizar_nombre(archivo.replace(".pdf", ".txt"))
        nombre_pdf = _normalizar_nombre(archivo)

        archivo_real_txt = archivos_en_disco.get(nombre_txt)
        archivo_real_pdf = archivos_en_disco.get(nombre_pdf)

        try:
            if archivo_real_txt:
                ruta = os.path.join(base, archivo_real_txt)
                with open(ruta, "r", encoding="utf-8") as f:
                    texto = f.read()
            elif archivo_real_pdf:
                ruta = os.path.join(base, archivo_real_pdf)
                reader = PdfReader(ruta)
                texto = "".join([p.extract_text() or "" for p in reader.pages])
            else:
                continue
            if texto.strip():
                resultado[nombre] = texto
        except Exception:
            pass
    return resultado

def buscar_fragmentos_relevantes(texto_doc: str, consulta: str,
                                  n_chunks: int = 6, chunk_size: int = 1200) -> str:
    """
    Divide el documento en fragmentos y devuelve los más relevantes
    según las palabras clave de la consulta.
    Evita mandar el PDF entero al modelo — solo los ~2700 chars más útiles.
    """
    # 1. Dividir respetando párrafos
    parrafos = texto_doc.split("\n")
    chunks, actual = [], ""
    for p in parrafos:
        if len(actual) + len(p) < chunk_size:
            actual += p + "\n"
        else:
            if actual.strip():
                chunks.append(actual.strip())
            actual = p + "\n"
    if actual.strip():
        chunks.append(actual.strip())

    if not chunks:
        return ""

    # 2. Palabras clave de la consulta (sin stopwords comunes)
    stopwords = {
        "de","la","el","en","que","y","a","los","las","un","con","del",
        "para","por","es","se","no","al","lo","una","como","o","sus",
        "si","sobre","pero","más","ya","puede","debe","tiene","hay",
    }
    palabras = set(consulta.lower().split()) - stopwords

    # 3. Puntuar cada chunk por coincidencias
    def puntuar(chunk):
        t = chunk.lower()
        return sum(1 for p in palabras if p in t)

    chunks_ordenados = sorted(chunks, key=puntuar, reverse=True)

    # 4. Devolver los N mejores (con al menos 1 punto para no mandar basura)
    top = [c for c in chunks_ordenados[:n_chunks] if puntuar(c) > 0]
    if not top:
        top = chunks_ordenados[:1]  # Al menos el primero si nada matchea

    return "\n\n---\n\n".join(top)

# ─────────────────────────────────────────────
# AVATARES SVG inline (base64 para usar en CSS)
# ─────────────────────────────────────────────

# Maestra primaria — delantal rosado, pelo recogido, sonrisa
SVG_PRIMARIO = """
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100" width="56" height="56">
  <!-- cuerpo delantal -->
  <ellipse cx="50" cy="78" rx="26" ry="22" fill="#f48fb1"/>
  <rect x="36" y="62" width="28" height="30" rx="6" fill="#f48fb1"/>
  <!-- cuello -->
  <rect x="44" y="56" width="12" height="12" rx="4" fill="#ffe0b2"/>
  <!-- cabeza -->
  <circle cx="50" cy="46" r="20" fill="#ffe0b2"/>
  <!-- pelo castaño recogido -->
  <ellipse cx="50" cy="30" rx="20" ry="12" fill="#6d4c41"/>
  <circle cx="50" cy="28" r="10" fill="#6d4c41"/>
  <circle cx="64" cy="33" r="7" fill="#6d4c41"/>
  <!-- rodete -->
  <circle cx="67" cy="28" r="6" fill="#6d4c41"/>
  <!-- ojos -->
  <circle cx="44" cy="46" r="3" fill="#fff"/>
  <circle cx="56" cy="46" r="3" fill="#fff"/>
  <circle cx="44.8" cy="46.5" r="1.6" fill="#333"/>
  <circle cx="56.8" cy="46.5" r="1.6" fill="#333"/>
  <!-- sonrisa -->
  <path d="M44 53 Q50 59 56 53" stroke="#c0392b" stroke-width="1.5" fill="none" stroke-linecap="round"/>
  <!-- mejillas -->
  <circle cx="41" cy="52" r="3" fill="#f8bbd0" opacity="0.7"/>
  <circle cx="59" cy="52" r="3" fill="#f8bbd0" opacity="0.7"/>
  <!-- manzana en mano -->
  <circle cx="76" cy="72" r="7" fill="#e53935"/>
  <path d="M76 65 Q78 62 80 64" stroke="#388e3c" stroke-width="1.5" fill="none"/>
</svg>
"""

# Tutor secundario — joven, camisa azul, corbata, pelo corto
SVG_SECUNDARIO = """
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100" width="56" height="56">
  <!-- cuerpo camisa -->
  <rect x="30" y="62" width="40" height="32" rx="8" fill="#1565c0"/>
  <!-- corbata -->
  <polygon points="50,64 53,72 50,80 47,72" fill="#e53935"/>
  <!-- cuello -->
  <rect x="43" y="56" width="14" height="12" rx="4" fill="#ffe0b2"/>
  <!-- solapa izq -->
  <polygon points="43,63 38,75 45,68" fill="#0d47a1"/>
  <!-- solapa der -->
  <polygon points="57,63 62,75 55,68" fill="#0d47a1"/>
  <!-- cabeza -->
  <circle cx="50" cy="45" r="20" fill="#ffe0b2"/>
  <!-- pelo corto oscuro -->
  <ellipse cx="50" cy="30" rx="20" ry="10" fill="#212121"/>
  <rect x="30" y="28" width="40" height="12" rx="6" fill="#212121"/>
  <!-- ojos -->
  <circle cx="44" cy="45" r="3" fill="#fff"/>
  <circle cx="56" cy="45" r="3" fill="#fff"/>
  <circle cx="44.8" cy="45.5" r="1.6" fill="#1a237e"/>
  <circle cx="56.8" cy="45.5" r="1.6" fill="#1a237e"/>
  <!-- sonrisa leve -->
  <path d="M45 52 Q50 57 55 52" stroke="#c0392b" stroke-width="1.5" fill="none" stroke-linecap="round"/>
  <!-- mejillas -->
  <circle cx="41" cy="50" r="3" fill="#ffccbc" opacity="0.6"/>
  <circle cx="59" cy="50" r="3" fill="#ffccbc" opacity="0.6"/>
</svg>
"""

# Profesor universitario — señor mayor, traje marrón, anteojos, barba canosa
SVG_UNIVERSIDAD = """
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 100 100" width="56" height="56">
  <!-- cuerpo traje marrón -->
  <rect x="28" y="62" width="44" height="32" rx="8" fill="#5d4037"/>
  <!-- camisa blanca interior -->
  <rect x="44" y="63" width="12" height="20" fill="#fafafa"/>
  <!-- corbata oscura -->
  <polygon points="50,64 52,71 50,78 48,71" fill="#3e2723"/>
  <!-- solapa izq -->
  <polygon points="44,63 30,80 42,70" fill="#4e342e"/>
  <!-- solapa der -->
  <polygon points="56,63 70,80 58,70" fill="#4e342e"/>
  <!-- cuello -->
  <rect x="43" y="55" width="14" height="12" rx="4" fill="#d7b896"/>
  <!-- cabeza -->
  <circle cx="50" cy="44" r="20" fill="#d7b896"/>
  <!-- pelo canoso corto -->
  <ellipse cx="50" cy="29" rx="20" ry="9" fill="#9e9e9e"/>
  <rect x="30" y="27" width="40" height="10" rx="5" fill="#9e9e9e"/>
  <!-- entradas -->
  <ellipse cx="33" cy="35" rx="5" ry="8" fill="#d7b896"/>
  <ellipse cx="67" cy="35" rx="5" ry="8" fill="#d7b896"/>
  <!-- barba canosa -->
  <ellipse cx="50" cy="60" rx="13" ry="7" fill="#bdbdbd" opacity="0.8"/>
  <!-- bigote -->
  <ellipse cx="50" cy="54" rx="8" ry="3" fill="#9e9e9e"/>
  <!-- ojos -->
  <circle cx="43" cy="44" r="3.5" fill="#fff"/>
  <circle cx="57" cy="44" r="3.5" fill="#fff"/>
  <circle cx="43.5" cy="44.5" r="1.8" fill="#4e342e"/>
  <circle cx="57.5" cy="44.5" r="1.8" fill="#4e342e"/>
  <!-- anteojos armazón -->
  <rect x="38" y="40" width="11" height="8" rx="3" fill="none" stroke="#5d4037" stroke-width="1.5"/>
  <rect x="51" y="40" width="11" height="8" rx="3" fill="none" stroke="#5d4037" stroke-width="1.5"/>
  <line x1="49" y1="44" x2="51" y2="44" stroke="#5d4037" stroke-width="1.5"/>
  <line x1="30" y1="43" x2="38" y2="43" stroke="#5d4037" stroke-width="1.5"/>
  <line x1="62" y1="43" x2="70" y2="43" stroke="#5d4037" stroke-width="1.5"/>
  <!-- expresión seria pero amable -->
  <path d="M45 52 Q50 55 55 52" stroke="#8d6e63" stroke-width="1.2" fill="none" stroke-linecap="round"/>
</svg>
"""

def svg_to_b64(svg_str):
    return base64.b64encode(svg_str.strip().encode()).decode()

AVATARES = {
    "Primario":    svg_to_b64(SVG_PRIMARIO),
    "Secundario":  svg_to_b64(SVG_SECUNDARIO),
    "Universidad": svg_to_b64(SVG_UNIVERSIDAD),
}

# ─────────────────────────────────────────────
# TEMAS DE COLOR POR NIVEL
# ─────────────────────────────────────────────
TEMAS = {
    "Primario": {
        "bg":           "linear-gradient(135deg, #fffde7 0%, #fff9c4 40%, #fff59d 70%, #fff176 100%)",
        "marco":        "#f9a825",
        "marco_glow":   "rgba(249,168,37,0.18)",
        "linea_asist":  "#ffe082",
        "linea_user":   "#fce4ec",
        "borde_asist":  "#f9a825",
        "borde_user":   "#e91e63",
        "input_borde":  "#f9a825",
        "titulo_color": "#e65100",
        "sidebar_bg":   "linear-gradient(180deg, #fb8c00 0%, #f57c00 50%, #e65100 100%)",
        "sidebar_borde":"#bf360c",
        "label_color":  "#fff9c4",
        "emoji_bar":    "🌈 ✏️ 🎨 📏 🌟",
        "spinner_msg":  "🍎 La seño está pensando...",
    },
    "Secundario": {
        "bg":           "linear-gradient(135deg, #e8f5e9 0%, #e3f2fd 40%, #ede7f6 70%, #e8eaf6 100%)",
        "marco":        "#5c6bc0",
        "marco_glow":   "rgba(92,107,192,0.15)",
        "linea_asist":  "#c5cae9",
        "linea_user":   "#fce4ec",
        "borde_asist":  "#3949ab",
        "borde_user":   "#e53935",
        "input_borde":  "#5c6bc0",
        "titulo_color": "#283593",
        "sidebar_bg":   "linear-gradient(180deg, #1a237e 0%, #283593 50%, #1a237e 100%)",
        "sidebar_borde":"#0d47a1",
        "label_color":  "#c5cae9",
        "emoji_bar":    "📱 🎧 ✏️ 💡 🚀",
        "spinner_msg":  "💬 Tu profe está respondiendo...",
    },
    "Universidad": {
        "bg":           "linear-gradient(135deg, #eceff1 0%, #e0e6ea 30%, #cfd8dc 60%, #b0bec5 100%)",
        "marco":        "#546e7a",
        "marco_glow":   "rgba(84,110,122,0.15)",
        "linea_asist":  "#b0bec5",
        "linea_user":   "#e8eaf6",
        "borde_asist":  "#37474f",
        "borde_user":   "#5c6bc0",
        "input_borde":  "#546e7a",
        "titulo_color": "#263238",
        "sidebar_bg":   "linear-gradient(180deg, #263238 0%, #37474f 50%, #263238 100%)",
        "sidebar_borde":"#102027",
        "label_color":  "#b0bec5",
        "emoji_bar":    "🔬 📐 🧮 📊 🎓",
        "spinner_msg":  "📖 El Dr. está elaborando la respuesta...",
    },
}

# ─────────────────────────────────────────────
# CONFIGURACIÓN DE PÁGINA (debe ir antes de todo widget)
# ─────────────────────────────────────────────
st.set_page_config(page_title="Tutor IA Multinivel", layout="centered", page_icon="🎓", initial_sidebar_state="expanded")

# Carga los PDFs de docs/ una sola vez (debe ir después de set_page_config)
DOCS_CARGADOS = cargar_documentos_referencia()

# Ocultar barra superior. En móvil ocultamos el sidebar nativo y mostramos un menú propio.
st.markdown("""
<style>
#MainMenu {visibility: hidden;}
header {visibility: hidden;}
footer {visibility: hidden;}

/* ── PC: sidebar fijo, visible, normal ── */
@media (min-width: 768px) {
    [data-testid="stSidebar"] {
        display: flex !important;
        visibility: visible !important;
        transform: none !important;
        position: relative !important;
    }
    [data-testid="collapsedControl"]      { display: none !important; }
    [data-testid="stSidebarCollapseButton"] { display: none !important; }
    #mobile-menu-bar { display: none !important; }
}

/* ── MÓVIL: ocultar sidebar nativo completamente ── */
@media (max-width: 767px) {
    [data-testid="stSidebar"]             { display: none !important; }
    [data-testid="collapsedControl"]      { display: none !important; }
    [data-testid="stSidebarCollapseButton"] { display: none !important; }
}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
defaults = {
    "autenticado": False,
    "chat_history": [],
    "contador": 0,
    "ultima_imagen_id": None,
    "descripcion_imagen": None,
    "nivel_actual": "Secundario",
    "nombre_alumno": "",
    "token_vence": "",
    "dias_restantes": 0,
    "ultimo_audio_id": None,
    "prompt_desde_audio": None,
    "ultima_respuesta_tts": None,
    "ultima_camara_id": None,
    "camara_b64_pendiente": None,
    "solicitar_desafio": False,
    "modo_docente": False,
    "modo_mixto": False,
    "modo_seleccionado": None,
    "sidebar_inicializado": False,
    "docs_ref_activos": {},   # { "NRA 2026": True/False, ... }
    "alerta_seguridad_activa": False,  # True si se disparó una alerta en esta sesión
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# Forzar sidebar abierto en la primera carga
if not st.session_state.get("sidebar_inicializado"):
    st.session_state["sidebar_inicializado"] = True
    st.rerun()

# ─────────────────────────────────────────────
# CSS DINÁMICO según nivel
# ─────────────────────────────────────────────
def inyectar_tema(nivel: str):
    t = TEMAS[nivel]
    av = AVATARES[nivel]
    st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Caveat:wght@400;600;700&family=Nunito:wght@400;600;700&display=swap');

.stApp {{
    background: {t['bg']};
    font-family: 'Nunito', sans-serif;
    transition: background 0.6s ease;
}}
.stApp h1 {{
    font-family: 'Caveat', cursive !important;
    font-size: 2.6rem !important;
    color: {t['titulo_color']} !important;
    text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    letter-spacing: 1px;
}}
[data-testid="stSidebar"] {{
    background: {t['sidebar_bg']} !important;
    border-right: 4px solid {t['sidebar_borde']};
    box-shadow: 4px 0 15px rgba(0,0,0,0.3);
}}
[data-testid="stSidebar"] * {{ color: #ecf0f1 !important; }}
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stFileUploader label {{
    color: {t['label_color']} !important;
    font-family: 'Caveat', cursive !important;
    font-size: 1.2rem !important;
    font-weight: 700 !important;
}}
[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p {{
    color: #ecf0f1 !important;
}}
/* ── FILE UPLOADER: botón y zona de arrastre bien visibles ── */
[data-testid="stSidebar"] [data-testid="stFileUploader"] section {{
    background: rgba(255,255,255,0.18) !important;
    border: 2px dashed rgba(255,255,255,0.7) !important;
    border-radius: 10px !important;
    padding: 10px !important;
}}
[data-testid="stSidebar"] [data-testid="stFileUploader"] section:hover {{
    background: rgba(255,255,255,0.28) !important;
    border-color: #fff !important;
}}
[data-testid="stSidebar"] [data-testid="stFileUploader"] button {{
    background: rgba(255,255,255,0.90) !important;
    color: #1a237e !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    padding: 6px 14px !important;
    box-shadow: 0 2px 6px rgba(0,0,0,0.25) !important;
}}
[data-testid="stSidebar"] [data-testid="stFileUploader"] button:hover {{
    background: #fff !important;
    transform: translateY(-1px) !important;
}}
[data-testid="stSidebar"] [data-testid="stFileUploader"] span,
[data-testid="stSidebar"] [data-testid="stFileUploader"] small,
[data-testid="stSidebar"] [data-testid="stFileUploader"] p {{
    color: rgba(255,255,255,0.92) !important;
    font-size: 0.82rem !important;
}}
[data-testid="stSidebar"] .stButton button {{
    background: linear-gradient(135deg, #e67e22, #d35400) !important;
    color: white !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Caveat', cursive !important;
    font-size: 1rem !important;
    font-weight: 700 !important;
    transition: all 0.2s ease !important;
    box-shadow: 0 3px 8px rgba(0,0,0,0.3) !important;
}}
[data-testid="stSidebar"] .stButton button:hover {{
    transform: translateY(-2px) !important;
    box-shadow: 0 5px 12px rgba(0,0,0,0.4) !important;
}}
/* ── MENSAJES ASISTENTE: fondo blanco + renglones encima ── */
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarAssistant"]) [data-testid="stChatMessageContent"] {{
    background-color: rgba(255,255,255,0.92) !important;
    background-image: repeating-linear-gradient(
        to bottom,
        transparent 0px, transparent 27px,
        {t['linea_asist']} 27px, {t['linea_asist']} 28px
    ) !important;
    background-size: 100% 28px !important;
    border-radius: 12px !important;
    border-left: 4px solid {t['borde_asist']} !important;
    padding: 12px 16px !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08) !important;
    font-family: 'Nunito', sans-serif !important;
    line-height: 28px !important;
}}
/* ── MENSAJES USUARIO: fondo blanco + renglones encima ── */
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) [data-testid="stChatMessageContent"] {{
    background-color: rgba(255,255,255,0.82) !important;
    background-image: repeating-linear-gradient(
        to bottom,
        transparent 0px, transparent 27px,
        {t['linea_user']} 27px, {t['linea_user']} 28px
    ) !important;
    background-size: 100% 28px !important;
    border-radius: 12px !important;
    border-left: 4px solid {t['borde_user']} !important;
    padding: 12px 16px !important;
    line-height: 28px !important;
}}
/* ── AVATAR: contenedor más grande y visible ── */
[data-testid="stChatMessageAvatarAssistant"] {{
    width: 52px !important;
    height: 52px !important;
    min-width: 52px !important;
    border-radius: 50% !important;
    border: 3px solid {t['borde_asist']} !important;
    background: white !important;
    overflow: hidden !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    box-shadow: 0 3px 10px rgba(0,0,0,0.2) !important;
    font-size: 2rem !important;
}}
[data-testid="stChatMessageAvatarUser"] {{
    width: 52px !important;
    height: 52px !important;
    min-width: 52px !important;
    border-radius: 50% !important;
    border: 3px solid {t['borde_user']} !important;
    box-shadow: 0 3px 10px rgba(0,0,0,0.15) !important;
    font-size: 2rem !important;
}}
/* Input */
[data-testid="stChatInput"] {{
    background: rgba(255,255,255,0.9) !important;
    border-radius: 16px !important;
    border: 2px solid {t['input_borde']} !important;
    box-shadow: 0 4px 12px {t['marco_glow']} !important;
}}
[data-testid="stChatInput"] textarea {{
    font-family: 'Caveat', cursive !important;
    font-size: 1.1rem !important;
    color: {t['titulo_color']} !important;
}}
/* Marco área principal */
.main .block-container {{
    background: rgba(255,255,255,0.6) !important;
    border-radius: 20px !important;
    border: 3px solid {t['marco']} !important;
    box-shadow:
        0 0 0 6px {t['marco_glow']},
        0 8px 32px rgba(0,0,0,0.12),
        inset 0 1px 0 rgba(255,255,255,0.8) !important;
    padding: 2rem 2.5rem !important;
    margin-top: 1rem !important;
}}
.pencil-bar {{
    display: flex; align-items: center; gap: 8px;
    margin-bottom: 0.5rem; font-size: 1.8rem; letter-spacing: 4px;
}}
[data-testid="stSidebar"] hr {{ border-color: rgba(255,255,255,0.2) !important; }}
[data-testid="stSidebar"] [data-testid="stAlert"] {{
    background: rgba(39,174,96,0.25) !important;
    border: 1px solid rgba(39,174,96,0.5) !important;
    border-radius: 8px !important;
}}
[data-testid="stExpander"] {{
    background: rgba(255,255,255,0.7) !important;
    border-radius: 10px !important;
    border: 1px solid {t['marco']} !important;
}}
.stDownloadButton button {{
    background: linear-gradient(135deg, #27ae60, #219a52) !important;
    color: white !important; border: none !important;
    border-radius: 8px !important;
    font-family: 'Caveat', cursive !important;
    font-size: 1rem !important; font-weight: 700 !important;
    width: 100% !important; margin-top: 8px !important;
    box-shadow: 0 3px 8px rgba(0,0,0,0.3) !important;
}}
.stDownloadButton button:hover {{
    background: linear-gradient(135deg, #2ecc71, #27ae60) !important;
    transform: translateY(-2px) !important;
}}
/* ── SELECTBOX NIVEL: bien visible ── */
[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] {{
    background: rgba(255,255,255,0.22) !important;
    border-radius: 10px !important;
    border: 2px solid rgba(255,255,255,0.7) !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.25) !important;
}}
[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"]:hover {{
    background: rgba(255,255,255,0.32) !important;
    border-color: #fff !important;
}}
[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] span {{
    color: #ffffff !important;
    font-family: 'Nunito', sans-serif !important;
    font-size: 1rem !important;
    font-weight: 700 !important;
}}
[data-testid="stSidebar"] .stSelectbox svg {{
    fill: #ffffff !important;
}}
/* ── AUDIO INPUT: micrófono resaltado ── */
[data-testid="stSidebar"] [data-testid="stAudioInput"] {{
    background: rgba(255,255,255,0.18) !important;
    border: 2px solid rgba(255,255,255,0.6) !important;
    border-radius: 12px !important;
    padding: 8px !important;
    box-shadow: 0 2px 10px rgba(0,0,0,0.25) !important;
}}
[data-testid="stSidebar"] [data-testid="stAudioInput"]:hover {{
    border-color: #fff !important;
    background: rgba(255,255,255,0.28) !important;
}}
[data-testid="stSidebar"] [data-testid="stAudioInput"] button {{
    background: linear-gradient(135deg, #e53935, #c62828) !important;
    border-radius: 50% !important;
    width: 42px !important;
    height: 42px !important;
    box-shadow: 0 3px 10px rgba(229,57,53,0.5) !important;
    border: none !important;
}}
[data-testid="stSidebar"] [data-testid="stAudioInput"] button:hover {{
    background: linear-gradient(135deg, #ff5252, #e53935) !important;
    transform: scale(1.08) !important;
    box-shadow: 0 4px 14px rgba(229,57,53,0.7) !important;
}}
[data-testid="stSidebar"] [data-testid="stAudioInput"] button svg {{
    fill: #ffffff !important;
}}
</style>
<div class="pencil-bar">{t['emoji_bar']}</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# SISTEMA DE TOKENS
# ─────────────────────────────────────────────
import hmac as _hmac
import hashlib as _hashlib
import base64 as _base64
from datetime import datetime as _datetime

def verificar_token(token: str) -> dict:
    """Verifica si el token es válido, auténtico y no venció."""
    try:
        clave = st.secrets.get("TOKEN_SECRET", "TutorIA_2025_ClaveSecreta_Cambiame")
        token_raw = _base64.urlsafe_b64decode(token.strip().encode()).decode()
        partes = token_raw.split("|")
        if len(partes) != 3:
            return {"valido": False, "motivo": "Token incorrecto"}
        nombre, vencimiento, firma_recibida = partes
        datos = f"{nombre}|{vencimiento}"
        firma_esperada = _hmac.new(
            clave.encode(), datos.encode(), _hashlib.sha256
        ).hexdigest()[:12]
        if not _hmac.compare_digest(firma_recibida, firma_esperada):
            return {"valido": False, "motivo": "Token inválido"}
        fecha_venc = _datetime.strptime(vencimiento, "%Y%m%d")
        if _datetime.now() > fecha_venc:
            dias_vencido = (_datetime.now() - fecha_venc).days
            return {"valido": False, "motivo": f"Tu acceso venció hace {dias_vencido} día(s). Renovalo contactando al profe."}
        dias_restantes = (fecha_venc - _datetime.now()).days
        return {"valido": True, "nombre": nombre,
                "vence": fecha_venc.strftime("%d/%m/%Y"), "dias_restantes": dias_restantes}
    except Exception:
        return {"valido": False, "motivo": "Token incorrecto. Verificá que lo copiaste bien."}

# ─────────────────────────────────────────────
# LOGIN CON TOKEN
# ─────────────────────────────────────────────
if not st.session_state.autenticado:
    inyectar_tema("Secundario")
    st.markdown("""
    <div style='text-align:center; padding: 2rem 0 1rem;'>
        <div style='font-size:4rem;'>🏫</div>
        <h1 style='font-family: Caveat, cursive; font-size:2.5rem; color:#283593;'>Aula Virtual IA</h1>
        <p style='font-family: Nunito, sans-serif; color:#555;'>
            Ingresá tu código de acceso para comenzar la clase.<br>
            <small>¿No tenés uno? Contactá al profe 
                <a href="https://wa.me/543584260947?text=Hola%2C%20te%20solicito%20un%20token%20de%20acceso.%20Mi%20nombre%20y%20apellido%20es%3A%20%5Bescribir%20nombre%20y%20apellido%5D%0ASoy%3A%20%5BAlumno%20%2F%20Docente%5D%0ANivel%3A%20%5BPrimario%20%2F%20Secundario%20%2F%20Universitario%5D" 
                   target="_blank" style="text-decoration:none; vertical-align:middle;">
                    <img src="https://upload.wikimedia.org/wikipedia/commons/6/6b/WhatsApp.svg" width="32" height="32" style="vertical-align:middle;" />
                </a>
            </small>
        </p>
    </div>""", unsafe_allow_html=True)
    st.markdown("""
<div style='background:rgba(255,255,255,0.12); border:1px solid rgba(255,255,255,0.3);
     border-radius:10px; padding:12px 16px; margin-bottom:16px;
     font-family:Nunito,sans-serif; font-size:0.78rem; color:#555; text-align:center;'>
    <b>⚖️ Términos de uso</b><br>
    Este servicio es una herramienta de apoyo educativo con inteligencia artificial.
    <b>No reemplaza a un docente ni constituye asesoramiento profesional.</b>
    Las respuestas pueden contener errores — verificá la información importante con tu docente.
    El uso del servicio implica la aceptación de estos términos.
    El acceso es personal e intransferible. El proveedor no se responsabiliza por
    decisiones tomadas en base al contenido generado por la IA.
</div>
""", unsafe_allow_html=True)
    token_input = st.text_input("🎟️ Código de acceso:", placeholder="Pegá tu token acá...").strip()
    col_a, col_b, col_c = st.columns([1,2,1])
    with col_b:
        if st.button("✏️ Entrar al Aula", use_container_width=True):
            if token_input:
                resultado = verificar_token(token_input)
                if resultado["valido"]:
                    st.session_state.autenticado      = True
                    st.session_state.nombre_alumno    = resultado["nombre"]
                    st.session_state.token_vence      = resultado["vence"]
                    st.session_state.dias_restantes   = resultado["dias_restantes"]
                    nombre_raw = resultado["nombre"]
                    st.session_state.modo_seleccionado = None
                    if nombre_raw.startswith("DOCENTE_ALUMNO_"):
                        st.session_state.modo_mixto    = True
                        st.session_state.modo_docente  = False
                        st.session_state.nombre_alumno = nombre_raw.replace("DOCENTE_ALUMNO_", "")
                    elif nombre_raw.startswith("DOCENTE_"):
                        st.session_state.modo_docente  = True
                        st.session_state.modo_mixto    = False
                        st.session_state.nombre_alumno = nombre_raw.replace("DOCENTE_", "")
                    else:
                        st.session_state.modo_docente  = False
                        st.session_state.modo_mixto    = False
                    st.rerun()
                else:
                    st.error(f"❌ {resultado['motivo']}")
            else:
                st.warning("Ingresá tu código de acceso.")
    st.stop()

# ─────────────────────────────────────────────
# MODELOS — API Key oculta en Streamlit Secrets
# ─────────────────────────────────────────────
os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]

MODEL_TEXT = "llama-3.3-70b-versatile"
VISION_MODELS = [
    "meta-llama/llama-4-scout-17b-16e-instruct",
    "meta-llama/llama-4-maverick-17b-128e-instruct",
]

try:
    llm_text = ChatGroq(model=MODEL_TEXT, temperature=0.1)
except Exception as e:
    st.error(f"Error de conexión: {e}")
    st.stop()

def get_vision_llm():
    for m in VISION_MODELS:
        try:
            return ChatGroq(model=m, temperature=0.1)
        except Exception:
            continue
    return None

def transcribir_audio(audio_bytes: bytes) -> str:
    """Transcribe audio usando Whisper de Groq."""
    try:
        import tempfile
        client = Groq(api_key=st.secrets["GROQ_API_KEY"])
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
            tmp.write(audio_bytes)
            tmp_path = tmp.name
        with open(tmp_path, "rb") as f:
            transcripcion = client.audio.transcriptions.create(
                model="whisper-large-v3",
                file=("audio.wav", f, "audio/wav"),
                language="es",
            )
        os.unlink(tmp_path)
        return transcripcion.text.strip()
    except Exception as e:
        return f"ERROR_AUDIO: {e}"

def texto_a_voz(texto: str):
    """Convierte texto a audio usando Groq TTS - Orpheus (modelo actual 2025)."""
    texto_corto = texto[:800] + ("..." if len(texto) > 800 else "")
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    try:
        response = client.audio.speech.create(
            model="canopylabs/orpheus-v1-english",
            voice="daniel",   # voz masculina clara
            input=texto_corto,
            response_format="wav",
        )
        return response.read(), None
    except Exception as e:
        return None, str(e)

def describir_imagen_automaticamente(img_b64: str) -> str:
    llm_vision = get_vision_llm()
    if llm_vision is None:
        return "No se pudo analizar la imagen (modelo de visión no disponible)."
    try:
        content = [
            {"type": "text", "text": (
                "Describí esta imagen de forma detallada y educativa. "
                "Indicá qué tipo de documento, ejercicio, problema o contenido contiene. "
                "Sé específico para que un tutor pueda responder preguntas sobre ella sin verla."
            )},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}}
        ]
        response = llm_vision.invoke([HumanMessage(content=content)])
        return response.content
    except Exception as e:
        return f"No se pudo analizar la imagen: {e}"

# ─────────────────────────────────────────────
# PRE-FILTRO DE SEGURIDAD (capa antes del LLM)
# Versión 2: multiidioma + leetspeak + fragmentación
# ─────────────────────────────────────────────
import re as _re
import unicodedata as _ud

def _normalizar_prefiltro(texto: str) -> str:
    """
    Normaliza el texto para resistir evasiones:
    - Quita tildes y diacríticos
    - Pasa a minúsculas
    - Reemplaza sustituciones leetspeak comunes (0→o, 1→i, 3→e, 4→a, @→a, $→s)
    - Elimina espacios y puntos entre letras de una misma palabra (s.u.i.c.i.d.i.o → suicidio)
    """
    # 1. Quitar tildes
    texto = _ud.normalize('NFD', texto).encode('ascii', 'ignore').decode('utf-8')
    # 2. Minúsculas
    texto = texto.lower()
    # 3. Leetspeak
    leet = {'0':'o','1':'i','3':'e','4':'a','@':'a','$':'s','5':'s','7':'t','!':'i'}
    texto = ''.join(leet.get(c, c) for c in texto)
    # 4. Puntos/guiones entre letras individuales (evasión tipo "s.u.i.c.i.d.i.o")
    texto = _re.sub(r'(?<=\b\w)[.\-_](?=\w\b)', '', texto)
    return texto

# ── Patrones en español ──
_PATRONES_ES = [
    # Suicidio y autolesión — palabras clave + combinaciones de pedido de métodos
    (r"(suicid\w*|autolesion\w*|quitarse\s+la\s+vida|hacerse\s+(dano|mal)|"
     r"cortarse\s+las?\s+venas?|tirarse?\s+(de|desde|por)\s+\w*\s*(altura|edificio|puente|balcon|ventana|piso)|"
     r"ahorcarse?|pastillas?\s+para\s+morir|sobredosis\s+intencional|"
     r"metodos?\s+(de|para)\s+(suicid\w*|morir|acabar|quitarse)|"
     r"como\s+(suicidarse|matarse|morir\s+sin\s+dolor|acabar\s+con\s+(todo|mi\s+vida)))",
     "crisis"),
    # Explosivos y armas
    (r"(tnt|explosiv\w*|detonar?\w*|fabricar\s+(bomba|explosivo|arma)|polvora|"
     r"petardo\s+casero|cable\s*(rojo|negro)\s*(explotar|detonar|encender|bomba)|"
     r"mezcla\s*(explosiva|peligrosa|quimica\s+peligrosa)|nitrato\s+de\s+amonio|"
     r"como\s+(hacer|fabricar|construir)\s+(una\s+)?(bomba|explosivo|arma\s+casera))",
     "peligro"),
    # Drogas
    (r"(como\s+(preparar|hacer|conseguir|fabricar|sintetizar)\s+"
     r"(droga\w*|cocaina|marihuana|pasta\s+base|fentanilo|heroina|metanfetamin\w*|crack|paco))",
     "drogas"),
]

# ── Patrones en inglés ──
_PATRONES_EN = [
    # Suicide and self-harm
    (r"(suicid\w*|self.?harm|cut\s+(my|your|the)\s+wrists?|"
     r"methods?\s+(of|for|to)\s+(suicide|killing\s+(myself|yourself)|ending\s+(my|your|a)\s+life)|"
     r"how\s+to\s+(kill\s+(my|your)self|commit\s+suicide|end\s+(my|your)\s+life|die\s+without\s+pain)|"
     r"ways?\s+to\s+(die|suicide|end\s+(my|your)\s+life)|"
     r"overdose\s+on\s+pills?|hanging\s+(my|your)self)",
     "crisis"),
    # Explosives and weapons
    (r"(tnt|explosiv\w*|detonat\w*|how\s+to\s+(make|build|create)\s+(a\s+)?(bomb|explosive|weapon)|"
     r"pipe\s+bomb|homemade\s+(bomb|explosive|weapon)|ammonium\s+nitrate\s+bomb)",
     "peligro"),
    # Drugs
    (r"(how\s+to\s+(make|cook|synthesize|produce)\s+(meth|heroin|cocaine|crack|fentanyl|crystal\s+meth))",
     "drogas"),
]

# ── Patrones en portugués ──
_PATRONES_PT = [
    (r"(suicid\w*|autolesao|como\s+(me\s+matar|suicidar|tirar\s+minha\s+vida)|"
     r"metodos?\s+(de\s+suicidio|para\s+morrer))",
     "crisis"),
    (r"(como\s+fazer\s+(uma\s+)?(bomba|explosivo)|tnt|explosiv\w*)",
     "peligro"),
    (r"(como\s+(fazer|preparar|sintetizar)\s+(droga\w*|cocaina|heroina|metanfetamin\w*))",
     "drogas"),
]

# ── Detector de fragmentación (acumulación de preguntas parciales peligrosas) ──
_FRAGMENTOS_PELIGROSOS = [
    r"(cuantas?\s+pastillas?|dosis\s+letal|dosis\s+maxima|que\s+pasa\s+si\s+tomo\s+muchas?)",
    r"(how\s+many\s+pills?|lethal\s+dose|what\s+happens\s+if\s+(i\s+take|you\s+take)\s+too\s+many)",
    r"(pastillas?\s+para\s+dormir|sleeping\s+pills?|somniferos?)",
]

_RESPUESTAS_SEGURIDAD = {
    "crisis": (
        "Noto que el tema que mencionás es muy sensible. "
        "Si vos o alguien que conocés está pasando por un momento difícil, "
        "lo más importante es hablar con alguien de confianza. "
        "En Argentina podés llamar al **135** (gratuito, las 24hs) — "
        "son personas capacitadas para escuchar y ayudar. 💙\n\n"
        "Yo solo puedo acompañarte con temas educativos, así que si querés, "
        "podemos retomar la clase cuando estés listo/a."
    ),
    "peligro": (
        "Eso que me pedís involucra materiales o procedimientos peligrosos "
        "y no puedo ayudarte con eso. Mi rol es acompañarte en el aprendizaje, "
        "no en actividades que puedan causar daño. 🚫\n\n"
        "Si tenés un proyecto escolar de ciencias, podemos buscar juntos "
        "una alternativa segura y aprobada por tu docente."
    ),
    "drogas": (
        "Ese tema está fuera de lo que puedo responder. "
        "Si tenés dudas sobre salud o prevención, hablá con un adulto de confianza "
        "o un profesional de la salud. 🙏\n\n"
        "¿Querés que sigamos con el contenido de la clase?"
    ),
    "fragmentacion": (
        "Esa combinación de preguntas me genera una alerta de seguridad y no puedo continuar "
        "por este camino. Si hay algo que te preocupa, hablá con un adulto de confianza "
        "o llamá al **135** (gratuito, 24hs). 💙\n\n"
        "Si es para un trabajo escolar, podemos reformular la consulta junto a tu docente."
    ),
}

# Historial de fragmentos sospechosos por sesión (se resetea con el chat)
_fragmentos_sesion: list[str] = []

def prefiltro_seguridad(texto: str, historial_reciente: list[str] | None = None, alerta_previa: bool = False) -> str | None:
    """
    Revisa el texto del usuario antes de mandarlo al LLM.
    Retorna una respuesta de seguridad si detecta contenido prohibido, o None si es seguro.
    Cubre: español, inglés, portugués, leetspeak, evasión con puntos, fragmentación
    y bloqueo por insistencia tras alerta previa en la misma sesión.
    """
    texto_norm = _normalizar_prefiltro(texto)

    # 0. Si ya se disparó una alerta antes en esta sesión, bloquear automáticamente
    #    cualquier mensaje que contenga palabras relacionadas con los temas sensibles
    if alerta_previa:
        _PALABRAS_SENSIBLES = (
            r"(pastillas?|medicamento\w*|dosis|somnifero\w*|sleeping|pills?|"
            r"suicid\w*|morir|muerte|matar\w*|daño|dolor|methods?|metodos?|"
            r"cuantas?|demasiadas?|how\s+many|too\s+many|enough\s+to)"
        )
        if _re.search(_PALABRAS_SENSIBLES, texto_norm):
            return _RESPUESTAS_SEGURIDAD["crisis"]

    # 1. Patrones directos en los tres idiomas
    for patrones in (_PATRONES_ES, _PATRONES_EN, _PATRONES_PT):
        for patron, categoria in patrones:
            if _re.search(patron, texto_norm):
                return _RESPUESTAS_SEGURIDAD[categoria]

    # 2. Detección de fragmentación acumulada
    fragmentos_en_texto = [p for p in _FRAGMENTOS_PELIGROSOS if _re.search(p, texto_norm)]
    if fragmentos_en_texto and historial_reciente:
        historial_norm = [_normalizar_prefiltro(h) for h in historial_reciente[-6:]]
        for frag in _FRAGMENTOS_PELIGROSOS:
            if any(_re.search(frag, h) for h in historial_norm):
                return _RESPUESTAS_SEGURIDAD["fragmentacion"]

    return None

# ─────────────────────────────────────────────
# AGENTE LANGGRAPH
# ─────────────────────────────────────────────
class AgentState(TypedDict):
    messages: List[BaseMessage]
    contexto_programa: str
    descripcion_imagen: str
    contador_pasos: int
    nivel_educativo: str

def tutor_node(state: AgentState):
    ultimo_msg = state['messages'][-1].content
    roles = {
        "Primario": (
            "Sos una maestra de primaria cariñosa y muy paciente. "
            "Usás palabras simples, analogías con juguetes o animales, frases cortas. "
            "Nunca usés fórmulas sin explicarlas con ejemplos del día a día. "
            "Celebrá cada avance del alumno con mucho entusiasmo y emojis."
        ),
        "Secundario": (
            "Sos un tutor de secundaria motivador y cercano. "
            "Usás lenguaje claro, ejemplos del mundo real y tecnología. "
            "Introducís terminología técnica pero siempre la explicás. "
            "Sos directo pero amable, como un compañero mayor que sabe mucho."
        ),
        "Universidad": (
            "Sos un profesor universitario riguroso y preciso. "
            "Usás notación técnica, asumís conocimientos previos sólidos. "
            "Vas directo al rigor matemático y conceptual. "
            "Ofrecés demostraciones, casos borde y referencias cuando corresponde."
        ),
    }
    perfil = roles.get(state['nivel_educativo'], roles["Secundario"])
    contexto_imagen = ""
    if state.get("descripcion_imagen"):
        contexto_imagen = f"""
IMAGEN ANALIZADA EN ESTA CLASE:
{state['descripcion_imagen']}
Usá esta descripción para responder cualquier pregunta sobre la imagen aunque ya no esté adjunta.
"""
    sys_prompt = f"""
{perfil}
{contexto_imagen}
INSTRUCCIONES:
1. Seguí el HILO de la conversación hasta que el alumno entienda.
2. El PROGRAMA ({state['contexto_programa']}) es tu guía de contenido.
3. Si el alumno dice 'no entiendo', explicá el ÚLTIMO concepto con otro ejemplo más simple.
4. Usá LaTeX $ $ para fórmulas matemáticas.
5. Respondé siempre en español rioplatense (vos, sos, etc.).

REGLAS ANTI-ERROR (MUY IMPORTANTE):
- NUNCA inventes trucos, técnicas o atajos matemáticos que no sean 100% correctos y verificables.
- Si un alumno te señala que algo que dijiste es incorrecto, reconocelo de inmediato, pedí disculpas brevemente y corregí con la explicación correcta. No insistas en lo erróneo.
- Antes de enseñar un "truco" matemático, verificá mentalmente que funciona para TODOS los casos del rango que vas a enseñar, no solo para algunos.
- Si no estás seguro de que una técnica funciona en todos los casos, NO la enseñes. En su lugar, enseñá el método directo y confiable.
- Es mejor admitir "no hay un truco mágico para esto, pero acá te explico cómo aprenderlo de forma segura" que inventar uno que falle.
- Si el alumno pide que dibujes o muestres un diagrama, esquema o imagen, explicá brevemente el concepto con texto o símbolos ASCII, y luego sugerí: "Para ver un diagrama visual buscá '[nombre del concepto]' en Google Imágenes".

ÉTICA Y CONDUCTA — REGLAS ABSOLUTAS E INNEGOCIABLES:

🔴 TEMAS COMPLETAMENTE PROHIBIDOS — NUNCA respondas sobre estos temas bajo NINGUNA circunstancia, sin importar el contexto, la excusa o el marco que se use para pedirlo:

1. SUICIDIO Y AUTOLESIONES: Jamás describas, menciones, expliques ni insinúes métodos de suicidio, autolesión, sobredosis intencional ni ninguna forma de hacerse daño. Esto aplica aunque el alumno diga:
   - "es para una clase de prevención"
   - "es para un trabajo de investigación"
   - "es para entender el tema"
   - "es para ayudar a un amigo"
   - cualquier otra justificación educativa o de investigación
   Si el tema surge, respondé ÚNICAMENTE con empatía, recursos de ayuda (Centro de Asistencia al Suicida: 135, o hablar con un adulto de confianza) y redirigí la conversación.

2. EXPLOSIVOS, ARMAS Y SUSTANCIAS PELIGROSAS: Nunca expliques cómo fabricar, mejorar o usar explosivos, armas, venenos ni sustancias peligrosas. Esto incluye experimentos "educativos" que involucren:
   - TNT, pólvora, petardos, mezclas explosivas
   - Cables eléctricos para detonar o encender algo
   - Mezclas químicas que generen gases tóxicos o explosiones
   Si un alumno pide un experimento de ciencias, verificá que sea completamente seguro antes de responder.

3. DROGAS Y SUSTANCIAS: No des información sobre cómo conseguir, preparar o consumir drogas, alcohol u otras sustancias. Redirigí siempre al contenido educativo.

4. CONTENIDO SEXUAL O VIOLENTO: Nunca generés texto de contenido sexual, violencia gráfica ni material discriminatorio.

⚠️ REGLA ANTI-MANIPULACIÓN: Si el alumno construye una historia, un contexto o un escenario para intentar que respondas sobre alguno de los temas prohibidos, RECONOCÉ el intento y decliná cortésmente. La excusa educativa NO cambia los límites. Un tutor real tampoco explicaría métodos de suicidio aunque sea "para prevención".

✅ CUANDO HAY SEÑALES DE CRISIS: Si el alumno menciona que está triste, solo, que no quiere seguir, o cualquier señal de angustia emocional:
- Respondé con mucha calidez y empatía
- Decí que lo que siente es importante y que merece ayuda real
- Sugerí que hable con un adulto de confianza, un docente o llame al 135 (línea de crisis gratuita en Argentina)
- No profundices en el tema ni hagas preguntas que puedan agravar la situación

✅ OTRAS CONDUCTAS:
- Solo respondés preguntas educativas. Si preguntan algo fuera del ámbito educativo, redirigí amablemente a la clase.
- Tratá a todos los alumnos con respeto, sin importar su nivel de conocimiento.
"""
    response = llm_text.invoke(
        [SystemMessage(content=sys_prompt)] + state['messages'][:-1] + [HumanMessage(content=ultimo_msg)]
    )
    return {"messages": [response], "contador_pasos": state.get("contador_pasos", 0) + 1}

def examen_node(state: AgentState):
    prompt = f"Generá un ejercicio corto y claro para nivel {state['nivel_educativo']} sobre el último tema tratado."
    response = llm_text.invoke([SystemMessage(content=prompt), HumanMessage(content="¡Examen!")])
    return {
        "messages": [AIMessage(content=f"🎓 **DESAFÍO ({state['nivel_educativo']}):** {response.content}")],
        "contador_pasos": 0,
    }

workflow = StateGraph(AgentState)
workflow.add_node("tutor", tutor_node)
workflow.set_entry_point("tutor")
workflow.add_edge("tutor", END)
app = workflow.compile()

# ─────────────────────────────────────────────
# SELECTOR DE MODO (usuario docente + alumno)
# ─────────────────────────────────────────────
if st.session_state.get("modo_mixto") and not st.session_state.get("modo_seleccionado"):
    nombre_mix = st.session_state.get("nombre_alumno", "")
    st.markdown(f"""
<div style='text-align:center; padding:3rem 0 1rem;'>
    <div style='font-size:3.5rem;'>🎓</div>
    <h1 style='font-family:Caveat,cursive; font-size:2.2rem; color:#283593;'>Bienvenido/a, {nombre_mix}</h1>
    <p style='font-family:Nunito,sans-serif; color:#555; font-size:1.05rem; margin-top:8px;'>
        ¿Con qué rol querés ingresar hoy?
    </p>
</div>""", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
<div style='background:rgba(40,53,147,0.08); border:2px solid rgba(40,53,147,0.2);
     border-radius:12px; padding:28px; text-align:center;'>
    <div style='font-size:2.5rem;'>🎒</div>
    <h3 style='font-family:Caveat,cursive; color:#283593; margin:10px 0 6px;'>Modo Alumno</h3>
    <p style='font-size:0.88rem; color:#666;'>Tutor IA disponible 24hs para tus consultas y estudio</p>
</div>""", unsafe_allow_html=True)
        if st.button("Entrar como Alumno", use_container_width=True, key="sel_alumno"):
            st.session_state.modo_seleccionado = "alumno"
            st.session_state.modo_docente = False
            st.rerun()
    with col2:
        st.markdown("""
<div style='background:rgba(39,174,96,0.08); border:2px solid rgba(39,174,96,0.2);
     border-radius:12px; padding:28px; text-align:center;'>
    <div style='font-size:2.5rem;'>👨‍🏫</div>
    <h3 style='font-family:Caveat,cursive; color:#1a7a4a; margin:10px 0 6px;'>Modo Docente</h3>
    <p style='font-size:0.88rem; color:#666;'>Asistente pedagógico para planificaciones y material didáctico</p>
</div>""", unsafe_allow_html=True)
        if st.button("Entrar como Docente", use_container_width=True, key="sel_docente"):
            st.session_state.modo_seleccionado = "docente"
            st.session_state.modo_docente = True
            st.rerun()
    st.stop()

# ─────────────────────────────────────────────
# MODO DOCENTE
# ─────────────────────────────────────────────
if st.session_state.get("modo_docente"):
    nombre_doc = st.session_state.nombre_alumno
    
    # Sidebar docente
    with st.sidebar:
        st.markdown("<div style='font-family:Caveat,cursive;font-size:1.4rem;color:#f0e68c;text-align:center;'>👨‍🏫 Asistente Docente</div>", unsafe_allow_html=True)
        st.markdown(f"""
<div style='background:rgba(39,174,96,0.25);border:1px solid rgba(39,174,96,0.5);
     border-radius:8px;padding:8px 12px;text-align:center;'>
  <div style='font-family:Caveat,cursive;font-size:1.1rem;color:#fff;'>✅ Prof. {nombre_doc}</div>
  <div style='font-size:0.75rem;color:rgba(255,255,255,0.75);'>
    Acceso hasta: {st.session_state.token_vence} · {st.session_state.dias_restantes}d restantes
  </div>
</div>""", unsafe_allow_html=True)

        st.divider()
        _h_list = ["Planificación de clase","Diseño de evaluación","Secuencia didáctica","Actividades para el aula","Adaptación para distintos niveles","Consulta pedagógica libre"]
        _h_def = st.session_state.get("_mob_herramienta", "Planificación de clase")
        _h_idx = _h_list.index(_h_def) if _h_def in _h_list else 0
        herramienta = st.selectbox("🛠️ ¿Qué necesitás?", _h_list, index=_h_idx)
        _n_list = ["Primario", "Secundario", "Universidad"]
        _n_def = st.session_state.get("_mob_nivel_doc", "Secundario")
        _n_idx = _n_list.index(_n_def) if _n_def in _n_list else 1
        nivel_doc = st.selectbox("📚 Nivel:", _n_list, index=_n_idx)
        materia_doc = st.text_input("📖 Materia:", placeholder="Ej: Matemáticas, Física...", value=st.session_state.get("_mob_materia_doc", ""))
        # Guardar en session_state para sincronizar con menú móvil
        st.session_state["_mob_herramienta"] = herramienta
        st.session_state["_mob_nivel_doc"]   = nivel_doc
        st.session_state["_mob_materia_doc"] = materia_doc

        # ── DOCUMENTOS DE REFERENCIA (solo aparecen si el PDF existe en docs/) ──
        if DOCS_CARGADOS:
            st.divider()
            st.markdown(
                "<div style='font-family:Caveat,cursive;font-size:1.1rem;"
                "font-weight:700;color:#f0e68c;'>📚 Reglamentos y marcos</div>",
                unsafe_allow_html=True,
            )
            st.caption("Activá para basar la respuesta en estos documentos")
            for nombre_doc, cfg in DOCS_CONFIG.items():
                if nombre_doc not in DOCS_CARGADOS:
                    continue  # el PDF no está en la carpeta, no mostramos el checkbox
                key_cb = f"cb_doc_{nombre_doc}"
                activo = st.checkbox(
                    f"{cfg['emoji']} {nombre_doc}",
                    value=st.session_state["docs_ref_activos"].get(nombre_doc, False),
                    key=key_cb,
                    help=cfg["descripcion"],
                )
                st.session_state["docs_ref_activos"][nombre_doc] = activo
                if activo:
                    st.markdown(
                        f"<div style='font-size:0.72rem;color:rgba(255,255,0,0.85);"
                        f"margin-top:-6px;margin-bottom:4px;'>✅ Activo</div>",
                        unsafe_allow_html=True,
                    )

        st.divider()
        if st.button("🗑️ Nueva consulta", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.alerta_seguridad_activa = False
            st.rerun()
        if st.session_state.get("modo_mixto"):
            if st.button("🔄 Cambiar a Alumno", use_container_width=True):
                st.session_state.modo_docente    = False
                st.session_state.modo_seleccionado = None
                st.session_state.chat_history    = []
                st.rerun()
        if st.button("🚪 Salir", use_container_width=True):
            for k, v in defaults.items():
                st.session_state[k] = v
            st.rerun()

        st.divider()
        # PDF docente
        st.markdown(
            "<div style='font-family:Caveat,cursive;font-size:1.1rem;font-weight:700;color:#f0e68c;'>"
            "📄 Cargar material</div>", unsafe_allow_html=True
        )
        pdf_docente = st.file_uploader("PDF (programa, material)", type="pdf", key="pdf_docente")
        img_docente = st.file_uploader("Imagen (foto, ejercicio)", type=["jpg","png","jpeg"], key="img_docente")

        if pdf_docente:
            from pypdf import PdfReader as _PdfReader
            texto_pdf_doc = "".join([p.extract_text() or "" for p in _PdfReader(pdf_docente).pages])
            st.session_state["contexto_docente_pdf"] = texto_pdf_doc
            st.success("✅ PDF cargado en memoria")

        if img_docente:
            img_id_doc = f"{img_docente.name}_{img_docente.size}"
            if img_id_doc != st.session_state.get("ultima_img_docente_id"):
                st.session_state.ultima_img_docente_id = img_id_doc
                st.session_state["desc_img_docente"] = None
                with st.spinner("🔍 Analizando imagen..."):
                    img_b64_doc = base64.b64encode(img_docente.read()).decode("utf-8")
                    img_docente.seek(0)
                    st.session_state["desc_img_docente"] = describir_imagen_automaticamente(img_b64_doc)
            if st.session_state.get("desc_img_docente"):
                st.success("✅ Imagen analizada")
                with st.expander("👁️ Ver descripción"):
                    st.write(st.session_state["desc_img_docente"])

        st.divider()
        # Descarga siempre visible
        if st.session_state.chat_history:
            # Generar Word
            doc = _Document()
            doc.core_properties.title = "Consulta Docente - Asistente IA"
            titulo = doc.add_heading("Consulta Docente — Asistente Pedagógico IA", level=1)
            titulo.runs[0].font.color.rgb = _RGBColor(0x28, 0x35, 0x93)
            doc.add_paragraph(f"Herramienta: {herramienta}  |  Nivel: {nivel_doc}  |  Materia: {materia_doc or 'General'}")
            doc.add_paragraph("")
            for m in st.session_state.chat_history:
                if isinstance(m, HumanMessage):
                    p = doc.add_paragraph()
                    run = p.add_run("👨‍🏫 DOCENTE:")
                    run.bold = True
                    run.font.size = _Pt(11)
                    doc.add_paragraph(m.content)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run("🤖 ASISTENTE:")
                    run.bold = True
                    run.font.color.rgb = _RGBColor(0x28, 0x35, 0x93)
                    run.font.size = _Pt(11)
                    doc.add_paragraph(m.content)
                doc.add_paragraph("")
            buf = _io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button(
                "📄 Descargar en Word",
                data=buf.getvalue(),
                file_name="consulta_docente.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.caption("📄 La descarga aparece luego de la primera respuesta")

    # Interfaz principal docente
    st.title("👨‍🏫 Asistente Pedagógico IA")
    st.markdown(f"""
<div style='background:rgba(39,174,96,0.1);border-left:4px solid #27ae60;
     border-radius:4px;padding:12px 16px;margin-bottom:20px;
     font-family:Nunito,sans-serif;font-size:0.92rem;color:#555;'>
    Modo docente activo · Herramienta: <b>{herramienta}</b> · Nivel: <b>{nivel_doc}</b>
</div>""", unsafe_allow_html=True)

    # Mostrar historial
    for m in st.session_state.chat_history:
        if isinstance(m, AIMessage):
            with st.chat_message("assistant", avatar="🤖"):
                st.markdown(m.content)
        else:
            with st.chat_message("user", avatar="👨‍🏫"):
                st.markdown(m.content)

    # Input docente
    prompt_doc = st.chat_input("✏️ Describí lo que necesitás...")
    if prompt_doc:
        new_msg = HumanMessage(content=prompt_doc)
        st.session_state.chat_history.append(new_msg)
        with st.chat_message("user", avatar="👨‍🏫"):
            st.markdown(prompt_doc)

        contexto_pdf_doc = st.session_state.get("contexto_docente_pdf", "")
        contexto_img_doc = st.session_state.get("desc_img_docente", "")
        contexto_extra = ""
        if contexto_pdf_doc:
            contexto_extra += f"\n\nMATERIAL PDF CARGADO POR EL DOCENTE:\n{contexto_pdf_doc[:3000]}"
        if contexto_img_doc:
            contexto_extra += f"\n\nIMAGEN ANALIZADA:\n{contexto_img_doc}"

        # ── DOCUMENTOS DE REFERENCIA ACTIVOS (RAG por fragmentos) ──
        docs_activos = st.session_state.get("docs_ref_activos", {})
        fragmentos_usados = []
        for nombre_ref, esta_activo in docs_activos.items():
            if esta_activo and nombre_ref in DOCS_CARGADOS:
                cfg_ref = DOCS_CONFIG[nombre_ref]
                fragmentos = buscar_fragmentos_relevantes(
                    DOCS_CARGADOS[nombre_ref], prompt_doc
                )
                if fragmentos:
                    fragmentos_usados.append(nombre_ref)
                    contexto_extra += f"""

══════════════════════════════════════════
{cfg_ref['emoji']} {nombre_ref.upper()} — {cfg_ref['descripcion']}
Fragmentos más relevantes para esta consulta:
══════════════════════════════════════════
{fragmentos}
══════════════════════════════════════════
"""

        # Instrucción de citado solo si hay docs activos
        instruccion_cita = ""
        if fragmentos_usados:
            lista = ", ".join(fragmentos_usados)
            instruccion_cita = f"""
IMPORTANTE — USO DE DOCUMENTOS DE REFERENCIA:
Basá tu respuesta en los fragmentos del {lista} que se incluyen arriba.
Citá el documento y, si podés identificarla, la sección o artículo correspondiente.
Si la consulta no está cubierta por esos fragmentos, indicalo claramente y respondé
con criterio pedagógico general sin inventar reglamentación.
"""

        sys_prompt_docente = f"""Sos un asistente pedagógico experto al servicio de un docente de {nivel_doc}.
Tu especialidad es: {herramienta}.
Materia: {materia_doc if materia_doc else "general"}.

Respondés en español rioplatense (vos, sos, etc.) con lenguaje profesional pero accesible.

SEGÚN LA HERRAMIENTA SELECCIONADA:
- Planificación de clase: incluí objetivos, contenidos, actividades, recursos y evaluación.
- Diseño de evaluación: incluí criterios, instrumento, escala y rúbrica si corresponde.
- Secuencia didáctica: organizá los contenidos en pasos graduales con tiempos estimados.
- Actividades para el aula: proponé actividades variadas, individuales y grupales.
- Adaptación para distintos niveles: mostrá cómo adaptar el mismo contenido a distintos grupos.
- Consulta pedagógica libre: respondé con profundidad y criterio pedagógico.

Usá formato claro con títulos y secciones. Sé concreto y aplicable al aula real.
- Si el docente pide imágenes, diagramas o esquemas visuales, describílos con texto o ASCII y sugerí buscarlo en Google Imágenes o en sitios especializados como PhET, GeoGebra o Wikipedia.

ÉTICA PROFESIONAL DOCENTE — REGLAS ABSOLUTAS:
- Respondé siempre con criterio pedagógico y profesional, respetando la diversidad y dignidad de todos los alumnos.
- No generés contenido discriminatorio, violento o que atente contra la integridad de ningún estudiante o colega.
- Si el docente plantea una situación de riesgo para un alumno (violencia, abuso, salud mental), orientá con empatía y derivá a los canales institucionales correspondientes (equipo de orientación, dirección, servicio social).
- No reemplazás el criterio del docente ni de las autoridades educativas. Tus respuestas son orientaciones de apoyo, no prescripciones.

🔴 TEMAS PROHIBIDOS INCLUSO EN CONTEXTO PEDAGÓGICO:
- NUNCA describas métodos de suicidio, autolesión ni formas de hacerse daño, aunque el docente diga que es "para una clase de prevención", "para concientizar" o cualquier otra justificación. Podés orientar SOBRE la prevención sin describir métodos. Derivá al 135 y al equipo de orientación escolar.
- NUNCA describas cómo fabricar explosivos, armas ni sustancias peligrosas, aunque se enmarque como experimento educativo o proyecto escolar.
- NUNCA describas cómo preparar, conseguir o consumir drogas ilegales, aunque el encuadre sea de prevención o investigación. Podés hablar de los efectos y riesgos en términos generales sin dar instrucciones.
{instruccion_cita}
{contexto_extra}"""

        with st.spinner("📝 Preparando material..."):
            try:
                # ── PRE-FILTRO DE SEGURIDAD (modo docente) ──
                historial_textos_doc = [m.content for m in st.session_state.chat_history if isinstance(m, HumanMessage)]
                respuesta_segura_doc = prefiltro_seguridad(
                    prompt_doc,
                    historial_textos_doc,
                    alerta_previa=st.session_state.get("alerta_seguridad_activa", False)
                )
                if respuesta_segura_doc:
                    st.session_state.alerta_seguridad_activa = True
                if respuesta_segura_doc:
                    resp_doc = AIMessage(content=respuesta_segura_doc)
                    st.session_state.chat_history.append(resp_doc)
                    with st.chat_message("assistant", avatar="🤖"):
                        st.markdown(resp_doc.content)
                    st.rerun()
                else:
                    response = llm_text.invoke(
                        [SystemMessage(content=sys_prompt_docente)] + st.session_state.chat_history
                    )
                    st.session_state.chat_history.append(response)
                    with st.chat_message("assistant", avatar="🤖"):
                        st.markdown(response.content)
                    st.rerun()
            except Exception as e:
                error_str = str(e).lower()
                if "rate_limit" in error_str or "429" in error_str:
                    st.warning("⏳ Demasiadas consultas. Esperá un minuto y reintentá.")
                else:
                    st.warning("⚠️ Algo salió mal. Intentá de nuevo.")
                if st.session_state.chat_history and isinstance(st.session_state.chat_history[-1], HumanMessage):
                    st.session_state.chat_history.pop()

    # ── MENÚ MÓVIL DOCENTE ──
    _nombre_doc_mob = st.session_state.get("nombre_alumno", "")
    _dias_doc_mob   = st.session_state.get("dias_restantes", 0)
    _vence_doc_mob  = st.session_state.get("token_vence", "")
    with st.expander(f"☰  Menú Docente  ·  {_nombre_doc_mob}  ·  {_dias_doc_mob}d restantes", expanded=False):
        st.markdown(f"""
        <div style='background:rgba(39,174,96,0.15); border:1px solid rgba(39,174,96,0.4);
             border-radius:8px; padding:8px 12px; text-align:center; font-size:0.85rem; margin-bottom:8px;'>
            👨‍🏫 <b>Prof. {_nombre_doc_mob}</b> · Acceso hasta: {_vence_doc_mob}
        </div>
        """, unsafe_allow_html=True)

        herramienta_mob = st.selectbox("🛠️ Herramienta:", [
            "Planificación de clase", "Diseño de evaluación", "Secuencia didáctica",
            "Actividades para el aula", "Adaptación para distintos niveles", "Consulta pedagógica libre",
        ], key="herramienta_mob")
        nivel_doc_mob   = st.selectbox("📚 Nivel:", ["Primario", "Secundario", "Universidad"], key="nivel_doc_mob")
        materia_doc_mob = st.text_input("📖 Materia:", placeholder="Ej: Matemáticas...", key="materia_mob")

        # Guardar selecciones para que el sidebar las tome en el próximo rerun
        st.session_state["_mob_herramienta"] = herramienta_mob
        st.session_state["_mob_nivel_doc"]   = nivel_doc_mob
        st.session_state["_mob_materia_doc"] = materia_doc_mob

        # ── DOCUMENTOS DE REFERENCIA (menú móvil) ──
        if DOCS_CARGADOS:
            st.markdown("---")
            st.markdown("**📚 Reglamentos y marcos**")
            st.caption("Activá para basar la respuesta en estos documentos")
            for nombre_doc_mob, cfg_mob in DOCS_CONFIG.items():
                if nombre_doc_mob not in DOCS_CARGADOS:
                    continue
                key_cb_mob = f"cb_doc_mob_{nombre_doc_mob}"
                activo_mob = st.checkbox(
                    f"{cfg_mob['emoji']} {nombre_doc_mob}",
                    value=st.session_state["docs_ref_activos"].get(nombre_doc_mob, False),
                    key=key_cb_mob,
                    help=cfg_mob["descripcion"],
                )
                st.session_state["docs_ref_activos"][nombre_doc_mob] = activo_mob

        st.markdown("---")
        if st.button("🗑️ Nueva consulta", key="doc_mob_reiniciar", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.alerta_seguridad_activa = False
            st.rerun()
        if st.session_state.get("modo_mixto"):
            if st.button("🔄 Cambiar a Alumno", key="doc_mob_alumno", use_container_width=True):
                st.session_state.modo_docente      = False
                st.session_state.modo_seleccionado = None
                st.session_state.chat_history      = []
                st.rerun()
        if st.button("🚪 Salir", key="doc_mob_salir", use_container_width=True):
            for k, v in defaults.items():
                st.session_state[k] = v
            st.rerun()

        st.markdown("---")
        st.markdown("**📄 Cargar material**")
        pdf_doc_mob = st.file_uploader("PDF (programa, material)", type="pdf", key="pdf_doc_mob")
        img_doc_mob = st.file_uploader("Imagen (foto, ejercicio)", type=["jpg","png","jpeg"], key="img_doc_mob")
        if pdf_doc_mob:
            from pypdf import PdfReader as _PdfReaderDocMob
            texto_pdf_doc_mob = "".join([p.extract_text() or "" for p in _PdfReaderDocMob(pdf_doc_mob).pages])
            st.session_state["contexto_docente_pdf"] = texto_pdf_doc_mob
            st.success("✅ PDF cargado")
        if img_doc_mob:
            img_id_doc_mob = f"{img_doc_mob.name}_{img_doc_mob.size}"
            if img_id_doc_mob != st.session_state.get("ultima_img_docente_id"):
                st.session_state.ultima_img_docente_id = img_id_doc_mob
                st.session_state["desc_img_docente"] = None
                with st.spinner("🔍 Analizando imagen..."):
                    img_b64_doc_mob = base64.b64encode(img_doc_mob.read()).decode("utf-8")
                    img_doc_mob.seek(0)
                    st.session_state["desc_img_docente"] = describir_imagen_automaticamente(img_b64_doc_mob)
            if st.session_state.get("desc_img_docente"):
                st.success("✅ Imagen analizada")
                with st.expander("👁️ Ver descripción"):
                    st.write(st.session_state["desc_img_docente"])

        if st.session_state.chat_history:
            st.markdown("---")
            doc_dm = _Document()
            doc_dm.core_properties.title = "Consulta Docente - Asistente IA"
            tit_dm = doc_dm.add_heading("Consulta Docente — Asistente Pedagógico IA", level=1)
            tit_dm.runs[0].font.color.rgb = _RGBColor(0x28, 0x35, 0x93)
            doc_dm.add_paragraph(f"Herramienta: {st.session_state.get('_mob_herramienta','General')}  |  Nivel: {st.session_state.get('_mob_nivel_doc','Secundario')}  |  Materia: {st.session_state.get('_mob_materia_doc','General')}")
            doc_dm.add_paragraph("")
            for m in st.session_state.chat_history:
                if isinstance(m, HumanMessage):
                    p = doc_dm.add_paragraph()
                    run = p.add_run("👨‍🏫 DOCENTE:")
                    run.bold = True
                    run.font.size = _Pt(11)
                    doc_dm.add_paragraph(m.content)
                else:
                    p = doc_dm.add_paragraph()
                    run = p.add_run("🤖 ASISTENTE:")
                    run.bold = True
                    run.font.color.rgb = _RGBColor(0x28, 0x35, 0x93)
                    run.font.size = _Pt(11)
                    doc_dm.add_paragraph(m.content)
                doc_dm.add_paragraph("")
            buf_dm = _io.BytesIO()
            doc_dm.save(buf_dm)
            buf_dm.seek(0)
            st.download_button(
                "📄 Descargar en Word",
                data=buf_dm.getvalue(),
                file_name="consulta_docente.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="mob_doc_download_word",
                use_container_width=True
            )
        else:
            st.caption("📄 La descarga aparece luego de la primera respuesta")

    st.stop()

# ─────────────────────────────────────────────
# INTERFAZ PRINCIPAL
# ─────────────────────────────────────────────

# Sidebar primero para leer el nivel antes de inyectar el tema
with st.sidebar:
    st.markdown("<div style='font-family: Caveat, cursive; font-size:1.4rem; color:#f0e68c; text-align:center;'>🏫 Aula Virtual</div>", unsafe_allow_html=True)
    # Info del alumno logueado
    nombre = st.session_state.get("nombre_alumno", "")
    vence  = st.session_state.get("token_vence", "")
    dias   = st.session_state.get("dias_restantes", 0)
    if nombre:
        st.markdown(f"""
<div style='background:rgba(39,174,96,0.25); border:1px solid rgba(39,174,96,0.5);
     border-radius:8px; padding:8px 12px; text-align:center;'>
  <div style='font-family:Caveat,cursive; font-size:1.1rem; color:#fff;'>✅ {nombre}</div>
  <div style='font-size:0.75rem; color:rgba(255,255,255,0.75);'>
    Acceso hasta: {vence} · {dias}d restantes
  </div>
</div>""", unsafe_allow_html=True)
    else:
        st.success("✅ Conectado")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("🗑️ Reiniciar", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.alerta_seguridad_activa = False
            st.session_state.contador = 0
            st.session_state.ultima_imagen_id = None
            st.session_state.descripcion_imagen = None
            st.rerun()
    with col2:
        if st.button("🚪 Salir", use_container_width=True):
            for k, v in defaults.items():
                st.session_state[k] = v
            st.session_state.nombre_alumno  = ""
            st.session_state.token_vence    = ""
            st.session_state.dias_restantes = 0
            st.rerun()
    if st.session_state.get("modo_mixto"):
        if st.button("🔄 Cambiar a Docente", use_container_width=True):
            st.session_state.modo_docente      = True
            st.session_state.modo_seleccionado = None
            st.session_state.chat_history      = []
            st.rerun()

    st.divider()
    if not st.session_state.get("modo_docente"):
        def _on_nivel_change():
            nuevo = st.session_state["nivel_edu_sidebar"]
            if nuevo != st.session_state.nivel_actual:
                st.session_state.nivel_actual = nuevo
                st.session_state.chat_history = []
                st.session_state.alerta_seguridad_activa = False
                st.session_state.contador = 0
        st.selectbox(
            "📚 Nivel del Alumno:",
            ["Primario", "Secundario", "Universidad"],
            index=["Primario","Secundario","Universidad"].index(st.session_state.nivel_actual),
            key="nivel_edu_sidebar",
            on_change=_on_nivel_change,
        )
    nivel_edu = st.session_state.nivel_actual

    st.divider()

    st.divider()
    st.markdown(
        "<div style='font-family:Caveat,cursive; font-size:1.15rem; font-weight:700; color:#f0e68c;'>"
        "🎙️ Consulta por voz</div>",
        unsafe_allow_html=True
    )
    st.caption("1️⃣ Grabá  ·  2️⃣ Detené  ·  3️⃣ Se envía solo")
    audio_input = st.audio_input(" ", key="audio_consulta", label_visibility="collapsed")
    if audio_input is not None:
        audio_bytes = audio_input.getvalue()
        audio_id = str(len(audio_bytes))
        if audio_id != st.session_state.get("ultimo_audio_id"):
            st.session_state.ultimo_audio_id = audio_id
            with st.spinner("🎙️ Transcribiendo..."):
                texto_transcripto = transcribir_audio(audio_bytes)
            if texto_transcripto.startswith("ERROR_AUDIO:"):
                st.warning("No se pudo transcribir. Escribí tu consulta.")
            else:
                st.session_state.prompt_desde_audio = texto_transcripto
                st.success(f'✅ "{texto_transcripto}"')
    st.divider()

    # ── CÁMARA ──
    st.markdown(
        "<div style='font-family:Caveat,cursive; font-size:1.15rem; font-weight:700; color:#f0e68c;'>"
        "📷 Foto con cámara</div>",
        unsafe_allow_html=True
    )
    activar_camara = st.toggle("📸 Activar cámara", value=False, key="toggle_camara")
    if activar_camara:
        st.caption("Encuadrá el ejercicio y presioná el botón de abajo ↓")
        camara_foto = st.camera_input("📸 Tomar foto", key="camara_ejercicio")
        if camara_foto is not None:
            cam_id = str(len(camara_foto.getvalue()))
            if cam_id != st.session_state.get("ultima_camara_id"):
                st.session_state.ultima_camara_id = cam_id
                # Guardamos la foto pero NO analizamos todavía
                cam_b64 = base64.b64encode(camara_foto.getvalue()).decode("utf-8")
                st.session_state.camara_b64_pendiente = cam_b64
                st.session_state.descripcion_imagen = None
                st.session_state.ultima_imagen_id = cam_id
            st.success("✅ Foto lista — hacé tu consulta y la analizaré")
    else:
        if st.session_state.get("ultima_camara_id"):
            st.caption("📷 Foto en memoria · activá para cambiarla")

    st.divider()

    # ── TTS: ESCUCHAR ÚLTIMA RESPUESTA ──
    st.markdown(
        "<div style='font-family:Caveat,cursive; font-size:1.15rem; font-weight:700; color:#f0e68c;'>"
        "🔊 Escuchar respuesta</div>",
        unsafe_allow_html=True
    )
    ultima_resp = st.session_state.get("ultima_respuesta_tts")
    if ultima_resp:
        st.caption("Presioná para escuchar la última respuesta del tutor")
        if st.button("▶️ Reproducir", use_container_width=True, key="btn_tts"):
            with st.spinner("🔊 Generando audio..."):
                audio_bytes_tts, error_tts = texto_a_voz(ultima_resp)
            if audio_bytes_tts:
                st.audio(audio_bytes_tts, format="audio/wav", autoplay=True)
            else:
                st.warning(f"No se pudo generar el audio. Error: {error_tts}")
    else:
        st.caption("Esperá la primera respuesta del tutor")

    st.divider()
    pdf_file  = st.file_uploader("📄 Programa (PDF)", type="pdf")
    img_file  = st.file_uploader("🖼️ Foto Ejercicio", type=["jpg","png","jpeg"])

    if img_file:
        imagen_id = f"{img_file.name}_{img_file.size}"
        if imagen_id != st.session_state.ultima_imagen_id:
            st.session_state.ultima_imagen_id = imagen_id
            st.session_state.descripcion_imagen = None
            with st.spinner("🔍 Analizando imagen..."):
                img_b64_temp = base64.b64encode(img_file.read()).decode('utf-8')
                img_file.seek(0)
                st.session_state.descripcion_imagen = describir_imagen_automaticamente(img_b64_temp)
        caption = "✅ Analizada y en memoria" if st.session_state.descripcion_imagen else "⏳ Analizando..."
        st.image(img_file, caption=caption, use_container_width=True)
        if st.session_state.descripcion_imagen:
            with st.expander("👁️ Ver descripción detectada"):
                st.write(st.session_state.descripcion_imagen)
    elif st.session_state.ultima_imagen_id is not None:
        st.session_state.ultima_imagen_id = None
        st.session_state.descripcion_imagen = None

    # ── BOTÓN DESAFÍO VOLUNTARIO ──
    st.divider()
    st.markdown(
        "<div style='font-family:Caveat,cursive; font-size:1.15rem; font-weight:700; color:#f0e68c;'>"
        "🎯 Evaluación voluntaria</div>",
        unsafe_allow_html=True
    )
    st.caption("Cuando te sentís listo, pedí un desafío sobre lo que estuvimos viendo.")
    if st.button("🎯 ¡Quiero ser evaluado!", use_container_width=True, key="btn_desafio"):
        st.session_state.solicitar_desafio = True
        st.rerun()

    if st.session_state.chat_history:
        st.divider()
        chat_text = "--- RESUMEN DE CLASE ---\n\n"
        for m in st.session_state.chat_history:
            autor = "ALUMNO" if isinstance(m, HumanMessage) else "PROFESOR"
            chat_text += f"[{autor}]: {m.content}\n\n"
        from docx import Document as _DocAlumno
        from docx.shared import Pt as _PtA, RGBColor as _RGBColorA
        import io as _ioA
        doc_a = _DocAlumno()
        doc_a.add_heading("Resumen de Clase — Aula Virtual IA", level=1)
        doc_a.add_paragraph(f"Alumno: {st.session_state.get('nombre_alumno','')}  |  Nivel: {nivel_edu}")
        doc_a.add_paragraph("")
        for m in st.session_state.chat_history:
            if isinstance(m, HumanMessage):
                p = doc_a.add_paragraph()
                run = p.add_run("👤 ALUMNO:")
                run.bold = True
                run.font.size = _PtA(11)
                doc_a.add_paragraph(m.content)
            else:
                p = doc_a.add_paragraph()
                run = p.add_run("🤖 TUTOR:")
                run.bold = True
                run.font.color.rgb = _RGBColorA(0x28, 0x35, 0x93)
                run.font.size = _PtA(11)
                doc_a.add_paragraph(m.content)
            doc_a.add_paragraph("")
        buf_a = _ioA.BytesIO()
        doc_a.save(buf_a)
        buf_a.seek(0)
        st.download_button(
            "📄 Descargar Clase en Word",
            data=buf_a.getvalue(),
            file_name="clase.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Inyectamos el tema DESPUÉS de leer el nivel del selectbox
inyectar_tema(nivel_edu)

# ── MENÚ MÓVIL: reemplaza el sidebar en pantallas pequeñas ──
# Solo para modo alumno — el modo docente tiene su propio expander
if not st.session_state.get("modo_docente"):

    _nombre_mob = st.session_state.get("nombre_alumno", "")
    _vence_mob  = st.session_state.get("token_vence", "")
    _dias_mob   = st.session_state.get("dias_restantes", 0)

    # Ocultar en PC, mostrar en móvil usando CSS sobre el expander
    st.markdown("""
    <style>
    /* En PC ocultamos el expander-menú (el sidebar ya tiene todo) */
    @media (min-width: 768px) {
        div[data-testid="stExpander"]:has(#mob-menu-anchor) {
            display: none !important;
        }
    }
    </style>
    <span id="mob-menu-anchor" style="display:none"></span>
    """, unsafe_allow_html=True)

    with st.expander(f"☰  Menú  ·  {_nombre_mob}  ·  {_dias_mob}d restantes", expanded=False):
        st.markdown(f"""
        <div style='background:rgba(39,174,96,0.15); border:1px solid rgba(39,174,96,0.4);
             border-radius:8px; padding:8px 12px; text-align:center;
             font-size:0.85rem; margin-bottom:8px;'>
            ✅ <b>{_nombre_mob}</b> · Acceso hasta: {_vence_mob}
        </div>
        """, unsafe_allow_html=True)

        def _on_nivel_mob_change():
            nuevo_mob = st.session_state["nivel_mob_select"]
            if nuevo_mob != st.session_state.nivel_actual:
                st.session_state.nivel_actual = nuevo_mob
                st.session_state.chat_history = []
                st.session_state.alerta_seguridad_activa = False
                st.session_state.contador = 0
        st.selectbox(
            "📚 Nivel:",
            ["Primario", "Secundario", "Universidad"],
            index=["Primario","Secundario","Universidad"].index(st.session_state.nivel_actual),
            key="nivel_mob_select",
            on_change=_on_nivel_mob_change,
        )

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            if st.button("🗑️ Reiniciar", key="mob_reiniciar", use_container_width=True):
                st.session_state.chat_history = []
                st.session_state.alerta_seguridad_activa = False
                st.session_state.contador = 0
                st.session_state.ultima_imagen_id = None
                st.session_state.descripcion_imagen = None
                st.rerun()
        with col_m2:
            if st.button("🚪 Salir", key="mob_salir", use_container_width=True):
                for k, v in defaults.items():
                    st.session_state[k] = v
                st.session_state.nombre_alumno  = ""
                st.session_state.token_vence    = ""
                st.session_state.dias_restantes = 0
                st.rerun()

        if st.session_state.get("modo_mixto"):
            if st.button("🔄 Cambiar a Docente", key="mob_docente", use_container_width=True):
                st.session_state.modo_docente      = True
                st.session_state.modo_seleccionado = None
                st.session_state.chat_history      = []
                st.rerun()

        st.markdown("---")
        st.markdown("**🎙️ Consulta por voz**")
        st.caption("1️⃣ Grabá  ·  2️⃣ Detené  ·  3️⃣ Se envía solo")
        audio_mob = st.audio_input(" ", key="audio_mob", label_visibility="collapsed")
        if audio_mob is not None:
            audio_bytes_mob = audio_mob.getvalue()
            audio_id_mob = str(len(audio_bytes_mob))
            if audio_id_mob != st.session_state.get("ultimo_audio_id"):
                st.session_state.ultimo_audio_id = audio_id_mob
                with st.spinner("🎙️ Transcribiendo..."):
                    texto_mob = transcribir_audio(audio_bytes_mob)
                if texto_mob.startswith("ERROR_AUDIO:"):
                    st.warning("No se pudo transcribir. Escribí tu consulta.")
                else:
                    st.session_state.prompt_desde_audio = texto_mob
                    st.success(f'✅ "{texto_mob}"')

        st.markdown("---")
        st.markdown("**🖼️ Subir imagen / foto**")
        img_mob = st.file_uploader("Foto ejercicio", type=["jpg","png","jpeg"], key="img_mob")
        if img_mob:
            imagen_id_mob = f"{img_mob.name}_{img_mob.size}"
            if imagen_id_mob != st.session_state.ultima_imagen_id:
                st.session_state.ultima_imagen_id = imagen_id_mob
                st.session_state.descripcion_imagen = None
                with st.spinner("🔍 Analizando imagen..."):
                    img_b64_mob = base64.b64encode(img_mob.read()).decode("utf-8")
                    img_mob.seek(0)
                    st.session_state.descripcion_imagen = describir_imagen_automaticamente(img_b64_mob)
            if st.session_state.descripcion_imagen:
                st.success("✅ Imagen analizada")

        st.markdown("**📷 Foto con cámara**")
        activar_camara_mob = st.toggle("📸 Activar cámara", value=False, key="toggle_camara_mob")
        if activar_camara_mob:
            st.caption("Encuadrá el ejercicio y presioná el botón de abajo ↓")
            camara_foto_mob = st.camera_input("📸 Tomar foto", key="camara_mob")
            if camara_foto_mob is not None:
                cam_id_mob = str(len(camara_foto_mob.getvalue()))
                if cam_id_mob != st.session_state.get("ultima_camara_id"):
                    st.session_state.ultima_camara_id = cam_id_mob
                    cam_b64_mob = base64.b64encode(camara_foto_mob.getvalue()).decode("utf-8")
                    st.session_state.camara_b64_pendiente = cam_b64_mob
                    st.session_state.descripcion_imagen = None
                    st.session_state.ultima_imagen_id = cam_id_mob
                st.success("✅ Foto lista — hacé tu consulta y la analizaré")
        else:
            if st.session_state.get("ultima_camara_id"):
                st.caption("📷 Foto en memoria · activá para cambiarla")

        st.markdown("---")
        st.markdown("**📄 Subir PDF del programa**")
        pdf_mob = st.file_uploader("PDF programa", type="pdf", key="pdf_mob")
        if pdf_mob:
            from pypdf import PdfReader as _PdfReaderMob
            contexto_mob = "".join([p.extract_text() or "" for p in _PdfReaderMob(pdf_mob).pages])
            st.session_state["contexto_pdf_mob"] = contexto_mob
            st.success("✅ PDF cargado")

        st.markdown("---")
        st.markdown("**🔊 Escuchar última respuesta**")
        ultima_resp_mob = st.session_state.get("ultima_respuesta_tts")
        if ultima_resp_mob:
            if st.button("▶️ Reproducir", use_container_width=True, key="btn_tts_mob"):
                with st.spinner("🔊 Generando audio..."):
                    audio_bytes_tts_mob, error_tts_mob = texto_a_voz(ultima_resp_mob)
                if audio_bytes_tts_mob:
                    st.audio(audio_bytes_tts_mob, format="audio/wav", autoplay=True)
                else:
                    st.warning(f"No se pudo generar el audio. Error: {error_tts_mob}")
        else:
            st.caption("Esperá la primera respuesta del tutor")

        st.markdown("---")
        if st.button("🎯 ¡Quiero ser evaluado!", key="mob_desafio", use_container_width=True):
            st.session_state.solicitar_desafio = True
            st.rerun()

        if st.session_state.chat_history:
            st.markdown("---")
            from docx import Document as _DocMob
            from docx.shared import Pt as _PtMob, RGBColor as _RGBMob
            import io as _ioMob
            doc_mob = _DocMob()
            doc_mob.add_heading("Resumen de Clase — Aula Virtual IA", level=1)
            doc_mob.add_paragraph(f"Alumno: {st.session_state.get('nombre_alumno','')}  |  Nivel: {nivel_edu}")
            doc_mob.add_paragraph("")
            for m in st.session_state.chat_history:
                if isinstance(m, HumanMessage):
                    p = doc_mob.add_paragraph()
                    run = p.add_run("👤 ALUMNO:")
                    run.bold = True
                    run.font.size = _PtMob(11)
                    doc_mob.add_paragraph(m.content)
                else:
                    p = doc_mob.add_paragraph()
                    run = p.add_run("🤖 TUTOR:")
                    run.bold = True
                    run.font.color.rgb = _RGBMob(0x28, 0x35, 0x93)
                    run.font.size = _PtMob(11)
                    doc_mob.add_paragraph(m.content)
                doc_mob.add_paragraph("")
            buf_mob = _ioMob.BytesIO()
            doc_mob.save(buf_mob)
            buf_mob.seek(0)
            st.download_button(
                "📄 Descargar Clase en Word",
                data=buf_mob.getvalue(),
                file_name="clase.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="mob_download_word",
                use_container_width=True
            )


# Títulos según nivel
titulos = {
    "Primario":    "👩‍🏫 Seño Virtual con Memoria",
    "Secundario":  "👨‍🏫 Tutor Agéntico con Memoria",
    "Universidad": "🎓 Profesor Virtual con Memoria",
}
st.title(titulos[nivel_edu])

# PDF
contexto = "General"
if pdf_file:
    contexto = "".join([p.extract_text() for p in PdfReader(pdf_file).pages])

# Chat — avatar dinámico según nivel
avatar_map = {
    "Primario":    "👩‍🏫",
    "Secundario":  "👨‍💼",
    "Universidad": "👨‍🔬",
}
avatar_asist = avatar_map[nivel_edu]
av_b64       = AVATARES[nivel_edu]
t_actual     = TEMAS[nivel_edu]

nombre_tutor = {
    "Primario":    "Seño Virtual 👩‍🏫",
    "Secundario":  "Tutor Virtual 👨‍💼",
    "Universidad": "Profesor Dr. 👨‍🔬",
}[nivel_edu]

saludo_tutor = {
    "Primario":    "¡Hola! Preguntame lo que quieras 🌟",
    "Secundario":  "Listo para ayudarte con lo que necesites 💡",
    "Universidad": "Proceda con su consulta académica 📚",
}[nivel_edu]

# Tarjeta de presentación del avatar grande
st.markdown(f"""
<div style="display:flex; align-items:center; gap:14px; margin-bottom:16px; padding:12px 16px;
     background:rgba(255,255,255,0.75); border-radius:16px;
     border-left: 5px solid {t_actual['borde_asist']};
     box-shadow: 0 3px 12px rgba(0,0,0,0.10);">
  <img src="data:image/svg+xml;base64,{av_b64}"
       style="width:80px; height:80px; border-radius:50%;
              border: 3px solid {t_actual['borde_asist']};
              background:white; box-shadow:0 4px 12px rgba(0,0,0,0.2);
              flex-shrink:0;" />
  <div>
    <div style="font-family:'Caveat',cursive; font-size:1.3rem;
                color:{t_actual['titulo_color']}; font-weight:700; margin-bottom:2px;">
      {nombre_tutor}
    </div>
    <div style="font-family:'Nunito',sans-serif; font-size:0.88rem; color:#666;">
      {saludo_tutor}
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

for m in st.session_state.chat_history:
    if isinstance(m, AIMessage):
        with st.chat_message("assistant", avatar=avatar_asist):
            st.markdown(m.content)
    else:
        with st.chat_message("user"):
            st.markdown(m.content)

spinner_msg = TEMAS[nivel_edu]["spinner_msg"]

with st.sidebar:
    pass  # asegura que el sidebar existe

st.markdown(
    "<div style='text-align:center; font-family:Nunito,sans-serif; font-size:0.72rem; "
    "color:#aaa; padding: 4px 0 8px;'>"
    "<b>Academia Particular IA</b> utiliza inteligencia artificial y puede cometer errores. "
    "Por favor, verificá las respuestas importantes con tu docente.</div>",
    unsafe_allow_html=True
)

# ── DESAFÍO VOLUNTARIO ──
if st.session_state.get("solicitar_desafio"):
    st.session_state.solicitar_desafio = False
    with st.spinner(spinner_msg):
        try:
            # Resumimos la conversación reciente para que el desafío sea sobre el tema actual
            temas_vistos = ""
            if st.session_state.chat_history:
                ultimos = st.session_state.chat_history[-6:]  # últimos 6 mensajes
                temas_vistos = "\n".join([
                    f"{'Alumno' if isinstance(m, HumanMessage) else 'Tutor'}: {m.content[:200]}"
                    for m in ultimos
                ])

            prompt_desafio = f"""Generá un ejercicio corto, claro y desafiante para nivel {nivel_edu} 
basado en los temas que se trataron en esta conversación reciente:

{temas_vistos if temas_vistos else "Tema general del nivel."}

El ejercicio debe tener 2 o 3 preguntas concretas. Luego de presentarlo, aclará que el alumno 
puede responder en el chat y que lo vas a corregir y ayudar si lo necesita."""

            response = llm_text.invoke([
                SystemMessage(content=prompt_desafio),
                HumanMessage(content="Generá el desafío ahora.")
            ])
            msg_desafio = AIMessage(content="🎯 **DESAFÍO (" + nivel_edu + "):**\n\n" + response.content)
            st.session_state.chat_history.append(msg_desafio)
            st.session_state.ultima_respuesta_tts = msg_desafio.content
            with st.chat_message("assistant", avatar=avatar_asist):
                st.markdown(msg_desafio.content)
        except Exception as e:
            error_str = str(e).lower()
            if "rate_limit" in error_str or "429" in error_str:
                st.warning("⏳ Demasiadas consultas. Esperá un minuto y volvé a intentar.")
            else:
                st.warning("⚠️ No se pudo generar el desafío. Intentá de nuevo en un momento.")

prompt_audio = st.session_state.get("prompt_desde_audio")
if prompt_audio:
    st.session_state.prompt_desde_audio = None

prompt_texto = st.chat_input("✏️ Escribí tu consulta acá...")
prompt = prompt_audio or prompt_texto

# Si hay foto de cámara pendiente de analizar, la analizamos ahora
if prompt and st.session_state.get("camara_b64_pendiente"):
    with st.spinner("🔍 Analizando tu foto..."):
        descripcion_cam = describir_imagen_automaticamente(st.session_state.camara_b64_pendiente)
        st.session_state.descripcion_imagen = descripcion_cam
        st.session_state.camara_b64_pendiente = None

if prompt:
    new_user_msg = HumanMessage(content=prompt)
    st.session_state.chat_history.append(new_user_msg)
    with st.chat_message("user"):
        if prompt_audio:
            st.markdown(f"🎙️ _{prompt}_")
        else:
            st.markdown(prompt)

    with st.spinner(spinner_msg):
        try:
            # ── PRE-FILTRO DE SEGURIDAD ──
            historial_textos = [m.content for m in st.session_state.chat_history if isinstance(m, HumanMessage)]
            respuesta_segura = prefiltro_seguridad(
                prompt,
                historial_textos,
                alerta_previa=st.session_state.get("alerta_seguridad_activa", False)
            )
            if respuesta_segura:
                st.session_state.alerta_seguridad_activa = True
            if respuesta_segura:
                resp_final = AIMessage(content=respuesta_segura)
                st.session_state.chat_history.append(resp_final)
                st.session_state.ultima_respuesta_tts = resp_final.content
                with st.chat_message("assistant", avatar=avatar_asist):
                    st.markdown(resp_final.content)
                st.rerun()
                st.stop()

            inputs = {
                "messages":          st.session_state.chat_history,
                "contexto_programa": contexto,
                "descripcion_imagen":st.session_state.descripcion_imagen,
                "contador_pasos":    st.session_state.contador,
                "nivel_educativo":   nivel_edu,
            }
            output     = app.invoke(inputs)
            resp_final = output["messages"][-1]
            st.session_state.contador = output.get("contador_pasos", 0)
            st.session_state.chat_history.append(resp_final)
            st.session_state.ultima_respuesta_tts = resp_final.content
            with st.chat_message("assistant", avatar=avatar_asist):
                st.markdown(resp_final.content)
            st.rerun()
        except Exception as e:
            error_str = str(e).lower()
            # Rate limit
            if "rate_limit" in error_str or "rate limit" in error_str or "429" in error_str:
                msgs = {
                    "Primario":    "⏳ ¡Uy! El profe está muy ocupado ahora. Esperá 1 minutito y volvé a preguntar. 😊",
                    "Secundario":  "⏳ Demasiadas consultas en este momento. Esperá un minuto y reintentá.",
                    "Universidad": "⏳ Límite de consultas alcanzado. Por favor aguarde 60 segundos antes de reintentar.",
                }
                st.warning(msgs.get(nivel_edu, msgs["Secundario"]))
            # Sin conexión / timeout
            elif "timeout" in error_str or "connection" in error_str or "network" in error_str:
                st.warning("🌐 Problema de conexión. Verificá tu internet y volvé a intentar.")
            # Token / auth
            elif "401" in error_str or "auth" in error_str or "api key" in error_str:
                st.error("🔑 Error de autenticación. Contactá al administrador.")
            # Cualquier otro error
            else:
                st.warning("⚠️ Algo salió mal. Esperá unos segundos y volvé a intentar. Si el problema persiste, usá el botón 🗑️ Reiniciar.")
            # Quitamos el último mensaje del historial para no dejar mensaje sin respuesta
            if st.session_state.chat_history and isinstance(st.session_state.chat_history[-1], HumanMessage):
                st.session_state.chat_history.pop()
